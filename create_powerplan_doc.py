import plan_csv_to_dict
import utils.general as general
from docx import Document
from docx.shared import Inches
import utils.content
import utils.styles
from datetime import datetime
from pathlib import Path

input_files = general.get_args()

powerplan_dict = plan_csv_to_dict.create_powerplan_dict(
    input_file=input_files.pathways_file,
    components_file=input_files.components_file,
    order_comments_file=input_files.order_comments_file,
)

DOMAIN = input_files.domain

SCRIPT_PATH = Path(__file__).parent


def remove_illegal_char_from_file_name(file_name: str) -> str:
    illegal = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for i in illegal:
        file_name = file_name.replace(i, "-")
    return file_name


for plan_dict in powerplan_dict.values():
    plan_description = plan_dict.get("POWERPLAN_DESCRIPTION")
    output_file_name = remove_illegal_char_from_file_name(plan_description)
    plan_display = plan_dict.get("POWERPLAN_DISPLAY")
    plan_type = plan_dict.get("PLAN_TYPE")
    plan_version = plan_dict.get("VERSION")
    begin_effective_date = plan_dict.get("BEG_EFFECTIVE_DT_TM")
    end_dttm = datetime.strptime(plan_dict.get("END_EFFECTIVE_DT_TM"),
                                 "%d-%b-%Y %H:%M")
    if end_dttm > datetime.now():
        end_effective_date = "Current"
    else:
        end_effective_date = plan_dict.get("END_EFFECTIVE_DT_TM")

    document = Document()

    document = utils.styles.add_styles_to_doc(document)

    document.sections[0].top_margin = Inches(0.53)
    document.sections[0].bottom_margin = Inches(0.53)
    document.sections[0].left_margin = Inches(0.53)
    document.sections[0].right_margin = Inches(0.53)

    document = utils.content.add_powerplan_rtf_header_to_doc(
        doc_obj=document,
        plan_desc=plan_description,
        plan_disp=plan_display,
        plan_type=plan_type,
        version=plan_version,
        begin_effective_dttm=begin_effective_date,
        end_effective_dttm=end_effective_date,
    )

    # Handling single-phase PowerPlans
    if not plan_dict.get("PHASES"):
        for comp_dict in sorted(
                plan_dict.get("COMPONENTS").values(),
                key=lambda x: x.get("SEQUENCE")
        ):
            p = document.add_paragraph("")
            # p.paragraph_format.left_indent = Inches(0.5)
            component = comp_dict.get('COMPONENT')
            comp_type = comp_dict.get("COMPONENT_TYPE")

            # Figure out the component offset quantity
            start_offset = comp_dict.get("START_OFFSET_QTY")
            zero_offset = comp_dict.get("ZERO_OFFSET_QTY")
            zero_time_ind = comp_dict.get("TIME_ZERO_IND")
            if start_offset:
                if start_offset[0] == "-":
                    offset = start_offset
                else:
                    offset = f"+{start_offset}"
            elif zero_offset:
                if start_offset[0] == "-":
                    offset = zero_offset
                else:
                    offset = f"+{zero_offset}"
            elif zero_time_ind:
                offset = "0 Hour"
            else:
                offset = ""

            # Figure out the component DOTs
            if comp_dict.get("DOT"):
                comp_dot_label = comp_dict.get("DOT_LABEL")[:-1]
                comp_dots = comp_dict.get("DOT").split(",")
                comp_dots_str = general.merge_list_of_numbers(comp_dots)
                comp_dot_disp = f"({comp_dot_label} {comp_dots_str})"
            else:
                comp_dot_disp = ""

            # Check box for orders
            if comp_type in ("Order", "Sub Phase", "Prescription"):
                if comp_dict.get("INCLUDE_IND"):
                    p.add_run(" ").add_picture("./tick_box.png")
                else:
                    p.add_run(" ").add_picture("./untick_box.png")

            # Component name, offset and DOTs
            if offset:
                p.add_run(f"\t{offset}").bold = True
                p.add_run(f" {component}")
            else:
                p.add_run(f"\t{component}")
            if comp_dot_disp:
                p.add_run(f" {comp_dot_disp}")

            # Abbreviations
            abbr = ""
            if comp_type == "Sub Phase":
                abbr = "SUB"
            elif comp_type == "Note":
                abbr = "NOTE"
            elif comp_type == "Prescription":
                abbr = "Rx"
            elif comp_dict.get("IV_COMPONENTS"):
                abbr = "IVS"
            if abbr:
                p.add_run(f"({abbr})*")

            # Order sentences
            if comp_dict.get("ORDER_SENTENCES"):
                for ord_sent_dict in comp_dict.get("ORDER_SENTENCES"):
                    if ord_sent_dict.get('ORDER_SENTENCE_DISPLAY_LINE'):
                        p2 = document.add_paragraph()
                        p2.paragraph_format.left_indent = Inches(1)
                        p2.add_run(
                            f"{ord_sent_dict.get('ORDER_SENTENCE_DISPLAY_LINE')}").italic = True
                    if ord_sent_dict.get('ORDER_COMMENTS'):
                        p3 = document.add_paragraph()
                        p3.paragraph_format.left_indent = Inches(1.5)
                        p3.add_run(
                            f"Comments: {ord_sent_dict.get('ORDER_COMMENTS')}").italic = True
            elif comp_dict.get("IV_COMPONENTS"):
                for row in comp_dict.get("IV_COMPONENTS"):
                    synonym = row.get("IV_SYNONYM")
                    order_sentence = row.get("ORDER_SENTENCE_DISPLAY_LINE")
                    order_comments = row.get("ORDER_COMMENTS")
                    document.add_paragraph(
                        f"{synonym}").paragraph_format.left_indent = Inches(1)
                    p4 = document.add_paragraph()
                    p4.paragraph_format.left_indent = Inches(1.5)
                    p4.add_run(f"{order_sentence}").italic = True

                    if order_comments:
                        p5 = document.add_paragraph()
                        p5.paragraph_format.left_indent = Inches(2)
                        p5.add_run(f"Comments: {order_comments}").italic = True

    # Multi-phase PowerPlans
    else:
        for phase_dict in sorted(plan_dict.get("PHASES").values(),
                                 key=lambda x: x.get("PHASE_SEQ")):
            phase_name = phase_dict.get("PHASE_DESCRIPTION")
            phase_duration = f"{phase_dict.get('PHASE_DURATION')}"
            if phase_dict.get("PHASE_DOT"):
                phase_dot_label = phase_dict.get("PHASE_DOT_LABEL")[:-1]
                phase_dots = phase_dict.get("PHASE_DOT").split(",")
                phase_dots_str = general.merge_list_of_numbers(phase_dots)
                phase_dot_disp = f"({phase_dot_label} {phase_dots_str})"
            else:
                phase_dot_disp = ""

            p1 = document.add_paragraph("")
            r1 = p1.add_run("")
            r1.add_break()
            p1.add_run(f"{phase_name} {phase_dot_disp}\t\t").bold = True
            if phase_duration:
                p1.add_run(f"Duration: {phase_duration}")

            # Start printing the components
            for comp_dict in sorted(
                    phase_dict.get("COMPONENTS").values(),
                    key=lambda x: x.get("SEQUENCE")
            ):
                p = document.add_paragraph("")
                # p.paragraph_format.left_indent = Inches(0.5)
                component = comp_dict.get('COMPONENT')
                comp_type = comp_dict.get("COMPONENT_TYPE")

                # Figure out the component offset quantity
                start_offset = comp_dict.get("START_OFFSET_QTY")
                zero_offset = comp_dict.get("ZERO_OFFSET_QTY")
                zero_time_ind = comp_dict.get("TIME_ZERO_IND")
                if start_offset:
                    if start_offset[0] == "-":
                        offset = start_offset
                    else:
                        offset = f"+{start_offset}"
                elif zero_offset:
                    if zero_offset[0] == "-":
                        offset = zero_offset
                    else:
                        offset = f"+{zero_offset}"
                elif zero_time_ind:
                    offset = "0 Hour"
                else:
                    offset = ""

                # Figure out the component DOTs
                if comp_dict.get("DOT"):
                    comp_dot_label = comp_dict.get("DOT_LABEL")[:-1]
                    comp_dots = comp_dict.get("DOT").split(",")
                    comp_dots_str = general.merge_list_of_numbers(
                        comp_dots)
                    comp_dot_disp = f"({comp_dot_label} {comp_dots_str})"
                else:
                    comp_dot_disp = ""

                # Check box for orders
                if comp_type in ("Order", "Sub Phase", "Prescription"):
                    if comp_dict.get("INCLUDE_IND"):
                        p.add_run(" ").add_picture("./tick_box.png")
                    else:
                        p.add_run(" ").add_picture("./untick_box.png")

                # Component name, offset and DOTs
                if offset:
                    p.add_run(f"\t{offset}").bold = True
                    p.add_run(f" {component}")
                else:
                    p.add_run(f"\t{component}")
                if comp_dot_disp:
                    p.add_run(f" {comp_dot_disp}")

                # Abbreviations
                abbr = ""
                if comp_type == "Sub Phase":
                    abbr = "SUB"
                elif comp_type == "Note":
                    abbr = "NOTE"
                elif comp_type == "Prescription":
                    abbr = "Rx"
                elif comp_dict.get("IV_COMPONENTS"):
                    abbr = "IVS"
                if abbr:
                    p.add_run(f"({abbr})*")

                # Order sentences
                if comp_dict.get("ORDER_SENTENCES"):
                    for ord_sent_dict in comp_dict.get("ORDER_SENTENCES"):

                        default_ind = False
                        if len(comp_dict.get("ORDER_SENTENCES")) > 1 and \
                            comp_dict.get("DEFAULT_OS_IND") == "1" and \
                            ord_sent_dict.get("ORDER_SENTENCE_SEQ") == "1":
                                default_ind = True
                        if ord_sent_dict.get('ORDER_SENTENCE_DISPLAY_LINE'):
                            p2 = document.add_paragraph()
                            p2.paragraph_format.left_indent = Inches(1)
                            p2.add_run(
                                f"{ord_sent_dict.get('ORDER_SENTENCE_DISPLAY_LINE')}").italic = True
                            if default_ind:
                                p2.add_run("(DEF)*")
                        if ord_sent_dict.get('ORDER_COMMENTS'):
                            p3 = document.add_paragraph()
                            p3.paragraph_format.left_indent = Inches(1.5)
                            p3.add_run(
                                f"Comments: {ord_sent_dict.get('ORDER_COMMENTS')}").italic = True
                elif comp_dict.get("IV_COMPONENTS"):
                    for row in comp_dict.get("IV_COMPONENTS"):
                        synonym = row.get("IV_SYNONYM")
                        order_sentence = row.get("ORDER_SENTENCE_DISPLAY_LINE")
                        order_comments = row.get("ORDER_COMMENTS")
                        document.add_paragraph(
                            f"{synonym}").paragraph_format.left_indent = Inches(
                            1)
                        p4 = document.add_paragraph()
                        p4.paragraph_format.left_indent = Inches(1.5)
                        p4.add_run(f"{order_sentence}").italic = True

                        if order_comments:
                            p5 = document.add_paragraph()
                            p5.paragraph_format.left_indent = Inches(2)
                            p5.add_run(
                                f"Comments: {order_comments}").italic = True

    utils.content.add_report_legend_to_document(document=document)

    utils.content.add_footer_to_document(document=document, domain=DOMAIN)
    try:
        document.save(Path(SCRIPT_PATH, "doc", f"{output_file_name}.docx"))
    except FileNotFoundError:
        print(output_file_name)
