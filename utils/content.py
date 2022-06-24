"""
utils/content.py
~~~~~~~~~~~~~~~~~~~~
Contains utilities dealing with adding content (as opposed to adding style
information)
"""
from datetime import datetime

import docx
from docx import Document
from docx.oxml import OxmlElement, ns

from utils.styles import add_bold_para_run_to_doc


def add_powerplan_rtf_header_to_doc(
    doc_obj: Document,
    plan_desc: str,
    plan_disp: str,
    plan_type: str,
    version: int = None,
    begin_effective_dttm: str = None,
    end_effective_dttm: str = None,
) -> Document:
    doc_obj = add_bold_para_run_to_doc(
        doc_obj=doc_obj, s=f"Unique Plan Description: {plan_desc}"
    )
    doc_obj = add_bold_para_run_to_doc(
        doc_obj=doc_obj, s=f"Plan Selection Display: {plan_disp}"
    )
    doc_obj = add_bold_para_run_to_doc(doc_obj=doc_obj, s=f"PlanType: {plan_type}")
    if version is not None:
        doc_obj = add_bold_para_run_to_doc(doc_obj=doc_obj, s=f"Version: {version}")
    if begin_effective_dttm is not None:
        doc_obj = add_bold_para_run_to_doc(
            doc_obj=doc_obj, s=f"Begin Effective Date: {begin_effective_dttm}"
        )
    if end_effective_dttm is not None:
        doc_obj = add_bold_para_run_to_doc(
            doc_obj=doc_obj, s=f"End Effective Date: {end_effective_dttm}"
        )
    return doc_obj



def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    page_num = create_element('w:instrText')
    create_attribute(page_num, 'xml:space', 'preserve')
    page_num.text = "PAGE"

    total_pages_num = create_element('w:instrText')
    create_attribute(total_pages_num, 'xml:space', 'preserve')
    total_pages_num.text = "NUMPAGES"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(page_num)
    run._r.append(fldChar2)


def add_total_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    page_num = create_element('w:instrText')
    create_attribute(page_num, 'xml:space', 'preserve')
    page_num.text = "NUMPAGES"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(page_num)
    run._r.append(fldChar2)


def add_report_legend_to_document(document: Document):
    p6 = document.add_paragraph()
    p6.add_run().add_break()
    p6.add_run("*Report Legend:").bold = True
    document.add_paragraph("DEF - This order sentence is the default for the selected order")
    document.add_paragraph("GOAL - This component is a goal")
    document.add_paragraph("IND - This component is an indicator")
    document.add_paragraph("INT - This component is an intervention")
    document.add_paragraph("IVS - This component is an IV Set")
    document.add_paragraph("NOTE - This component is a note")
    document.add_paragraph("Rx - This component is a prescription")
    document.add_paragraph("SUB - This component is a subphase")


def add_footer_to_document(document: Document, domain: str = ""):
    document.sections[0].footer.paragraphs[0].add_run(
        f"Printed on: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
    document.sections[0].footer.paragraphs[0].add_run("\t")
    document.sections[0].footer.paragraphs[0].add_run("Page ")
    add_page_number(
        document.sections[0].footer.paragraphs[0].add_run())
    document.sections[0].footer.paragraphs[0].add_run("   of ")
    add_total_page_number(
        document.sections[0].footer.paragraphs[0].add_run())
    document.sections[0].footer.paragraphs[0].add_run(f"\t Domain: {domain}")